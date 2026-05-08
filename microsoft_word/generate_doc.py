from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
BLACK      = RGBColor(0x0D, 0x0D, 0x0D)
ACCENT     = RGBColor(0x1A, 0x1A, 0x2E)   # deep navy
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)
MID_GREY   = RGBColor(0x88, 0x88, 0x88)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper functions ──────────────────────────────────────────────────────────
def set_font(run, size=11, bold=False, colour=BLACK, font_name="Calibri"):
    run.font.name        = font_name
    run.font.size        = Pt(size)
    run.font.bold        = bold
    run.font.color.rgb   = colour

def add_paragraph(doc, text="", size=11, bold=False, colour=BLACK,
                  align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=6,
                  font_name="Calibri"):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, size, bold, colour, font_name)
    return p

def add_heading(doc, text, level_size=20, colour=ACCENT, space_before=14, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, level_size, bold=True, colour=colour)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A1A2E')
    pBdr.append(bottom)
    pPr.append(pBdr)

def shade_paragraph(paragraph, hex_fill="1A1A2E"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
#  HEADER BANNER
# ══════════════════════════════════════════════════════════════════════════════
banner = doc.add_paragraph()
banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
banner.paragraph_format.space_before = Pt(0)
banner.paragraph_format.space_after  = Pt(0)
shade_paragraph(banner, "1A1A2E")
run = banner.add_run("  IAN KINOTI  ·  CREATIVE DEVELOPER  ·  KENYA  ")
set_font(run, size=10, bold=True, colour=WHITE, font_name="Calibri")

sub_banner = doc.add_paragraph()
sub_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_banner.paragraph_format.space_before = Pt(0)
sub_banner.paragraph_format.space_after  = Pt(10)
shade_paragraph(sub_banner, "1A1A2E")
run2 = sub_banner.add_run("kicoinkinoti@gmail.com  ·  https://my-portfolio-ten-jet-99.vercel.app")
set_font(run2, size=9, colour=RGBColor(0xCC, 0xCC, 0xCC), font_name="Calibri")

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — WELCOME LETTER
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "01  |  WELCOME LETTER", 16)
add_divider(doc)

add_paragraph(doc, f"Date: {datetime.date.today().strftime('%B %d, %Y')}", size=10, colour=MID_GREY, space_before=8)
add_paragraph(doc, "To: [CLIENT_FULL_NAME]", size=11, bold=True, space_before=10)
add_paragraph(doc, "[CLIENT_COMPANY / BUSINESS NAME]", size=11)
add_paragraph(doc, "[CLIENT_EMAIL]  ·  [CLIENT_PHONE]", size=10, colour=MID_GREY)

add_paragraph(doc, "", space_after=4)

add_paragraph(doc,
    "Dear [CLIENT_FIRST_NAME],",
    size=11, space_before=10, space_after=6)

add_paragraph(doc,
    "Thank you for choosing to work with me. I'm genuinely excited about what we're going to build together. "
    "This document serves as your official welcome package — it outlines the scope of our engagement, "
    "the investment required, and the terms that protect both of us throughout the project.",
    size=11, space_after=6)

add_paragraph(doc,
    "My commitment to you is simple: I will deliver a fast, modern, and polished digital product "
    "that reflects your brand and serves your goals. Every decision — from architecture to animation — "
    "will be made with intention and care.",
    size=11, space_after=6)

add_paragraph(doc,
    "Please read through this document carefully. If you have any questions before signing, "
    "do not hesitate to reach out at kicoinkinoti@gmail.com.",
    size=11, space_after=6)

add_paragraph(doc, "Warm regards,", size=11, space_before=10, space_after=2)
add_paragraph(doc, "Ian Kinoti", size=12, bold=True, space_after=2)
add_paragraph(doc, "Creative Developer  ·  Ian Creative", size=10, colour=MID_GREY)
add_paragraph(doc, "https://my-portfolio-ten-jet-99.vercel.app", size=10, colour=ACCENT)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — INVOICE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

add_heading(doc, "02  |  INVOICE", 16)
add_divider(doc)

# Meta table
meta = doc.add_table(rows=4, cols=2)
meta.style = 'Table Grid'
meta.alignment = WD_TABLE_ALIGNMENT.LEFT

meta_data = [
    ("Invoice No.",       "INV-[INVOICE_NUMBER]"),
    ("Invoice Date:",     datetime.date.today().strftime('%B %d, %Y')),
    ("Due Date:",         "[PAYMENT_DUE_DATE]"),
    ("Project Title:",    "[PROJECT_TITLE / WEBSITE TYPE]"),
]
for i, (label, value) in enumerate(meta_data):
    row = meta.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    for j, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.bold = (j == 0)

add_paragraph(doc, "", space_after=10)

# Billing parties table
add_paragraph(doc, "BILLED BY", size=9, bold=True, colour=MID_GREY, space_before=6, space_after=2)
add_paragraph(doc, "Ian Kinoti", size=11, bold=True, space_after=0)
add_paragraph(doc, "Creative Developer  ·  Nairobi, Kenya", size=10, space_after=0)
add_paragraph(doc, "kicoinkinoti@gmail.com", size=10, colour=ACCENT, space_after=10)

add_paragraph(doc, "BILLED TO", size=9, bold=True, colour=MID_GREY, space_before=6, space_after=2)
add_paragraph(doc, "[CLIENT_FULL_NAME]", size=11, bold=True, space_after=0)
add_paragraph(doc, "[CLIENT_COMPANY / BUSINESS NAME]", size=10, space_after=0)
add_paragraph(doc, "[CLIENT_EMAIL]  ·  [CLIENT_PHONE]", size=10, colour=MID_GREY, space_after=10)

# Line items table
add_paragraph(doc, "SERVICES", size=9, bold=True, colour=MID_GREY, space_before=6, space_after=4)

items_table = doc.add_table(rows=1, cols=4)
items_table.style = 'Table Grid'
items_table.alignment = WD_TABLE_ALIGNMENT.LEFT

hdr_cells = items_table.rows[0].cells
headers = ["#", "Description", "Qty", "Amount (KES)"]
col_widths = [Cm(1), Cm(9.5), Cm(2), Cm(4)]

for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
    cell.width = col_widths[i]
    para = cell.paragraphs[0]
    run = para.add_run(hdr)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = WHITE
    shade_cell(cell, "1A1A2E") if False else None   # handled below

    # shade header row
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A1A2E')
    tcPr.append(shd)
    run.font.color.rgb = WHITE

line_items = [
    ("1", "[SERVICE / DELIVERABLE 1 — e.g. Website Design & Development]", "1", "[AMOUNT]"),
    ("2", "[SERVICE / DELIVERABLE 2 — e.g. CMS Integration]",              "1", "[AMOUNT]"),
    ("3", "[SERVICE / DELIVERABLE 3 — e.g. Mobile Responsiveness]",        "1", "Included"),
    ("4", "[SERVICE / DELIVERABLE 4 — e.g. SEO Setup]",                    "1", "Included"),
]

for item in line_items:
    row_cells = items_table.add_row().cells
    for i, (cell, val) in enumerate(zip(row_cells, item)):
        cell.width = col_widths[i]
        para = cell.paragraphs[0]
        run = para.add_run(val)
        run.font.name = "Calibri"
        run.font.size = Pt(10)

# Totals
add_paragraph(doc, "", space_after=4)
totals_table = doc.add_table(rows=3, cols=2)
totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT

totals_data = [
    ("Subtotal:",  "KES [SUBTOTAL]"),
    ("Deposit (50% upfront):", "KES [DEPOSIT_AMOUNT]"),
    ("Balance Due on Delivery:", "KES [BALANCE_AMOUNT]"),
]
for i, (label, val) in enumerate(totals_data):
    row = totals_table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = val
    for j, cell in enumerate(row.cells):
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in para.runs:
                run.font.name  = "Calibri"
                run.font.size  = Pt(10)
                run.font.bold  = (i == 2)
                if i == 2:
                    run.font.color.rgb = ACCENT

add_paragraph(doc, "Payment via M-Pesa / Bank Transfer. Details provided separately.", size=9, colour=MID_GREY, space_before=8)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — CLIENT AGREEMENT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()

add_heading(doc, "03  |  CLIENT AGREEMENT", 16)
add_divider(doc)

add_paragraph(doc,
    "This Agreement is entered into as of the date signed below, between Ian Kinoti "
    "(\"Developer\") and [CLIENT_FULL_NAME] / [CLIENT_COMPANY] (\"Client\").",
    size=11, space_before=10, space_after=8)

# ── Clauses ──
clauses = [
    ("1. SCOPE OF WORK",
     "The Developer agrees to design and develop [PROJECT_TITLE] as described in the agreed project brief. "
     "Any features, pages, or functionality not explicitly stated in the brief are outside the scope of this agreement "
     "and will be quoted separately."),

    ("2. TIMELINE",
     "The estimated delivery timeline is [PROJECT_TIMELINE]. This timeline is contingent on the Client providing "
     "all required content, assets, and feedback within [FEEDBACK_TURNAROUND_DAYS] business days of each request. "
     "Delays caused by the Client will extend the delivery date accordingly."),

    ("3. PAYMENT TERMS",
     "A non-refundable deposit of 50% of the total project fee is due before work commences. "
     "The remaining balance is due upon project completion and before the final files or live deployment are handed over. "
     "Failure to pay the balance within 7 days of delivery notice may result in work being withheld."),

    ("4. REVISIONS",
     "This agreement includes [NUMBER_OF_REVISIONS] rounds of revisions. A revision is defined as minor "
     "adjustments to existing content or design — not a change of direction or new feature requests. "
     "Additional revisions beyond the included rounds will be billed at KES [HOURLY_RATE] per hour."),

    ("5. INTELLECTUAL PROPERTY",
     "Upon receipt of full payment, the Client receives full ownership of the final deliverables. "
     "The Developer retains the right to display the project in their portfolio and promotional materials "
     "unless the Client requests otherwise in writing."),

    ("6. CONFIDENTIALITY",
     "Both parties agree to keep all project-related information, business data, and proprietary content "
     "confidential and not to disclose it to third parties without prior written consent."),

    ("7. CLIENT SELF-MAINTENANCE & LIABILITY DISCLAIMER",
     "The Developer delivers a fully tested, functional product at the time of handover. "
     "Once the Client takes ownership of the codebase, hosting environment, or any aspect of the website, "
     "the Developer is no longer liable for any crashes, data loss, security breaches, broken functionality, "
     "or any damages — direct or indirect — arising from the Client's modifications, third-party plugin updates, "
     "hosting issues, or any changes made without the Developer's involvement.\n\n"
     "The Client acknowledges and accepts full responsibility for the website upon handover. "
     "Any post-handover technical support, maintenance, updates, or repairs will require a separate "
     "Maintenance Agreement and will be billed at the Developer's current rates."),

    ("8. HOSTING & DOMAIN",
     "Unless explicitly included in the project scope, the Client is responsible for procuring and maintaining "
     "their own hosting and domain. The Developer may assist with setup but is not responsible for ongoing "
     "hosting performance, renewal, or downtime caused by the Client's hosting provider."),

    ("9. TERMINATION",
     "Either party may terminate this agreement with written notice. If the Client terminates the project "
     "after work has commenced, the deposit is non-refundable and any work completed beyond the deposit value "
     "will be invoiced and due immediately. If the Developer terminates, a pro-rata refund of the deposit "
     "will be issued for work not yet performed."),

    ("10. GOVERNING LAW",
     "This agreement shall be governed by the laws of Kenya. Any disputes shall first be attempted to be "
     "resolved amicably. If unresolved, disputes shall be referred to the relevant jurisdiction in Nairobi, Kenya."),
]

for title, body in clauses:
    add_paragraph(doc, title, size=11, bold=True, colour=ACCENT, space_before=12, space_after=3)
    add_paragraph(doc, body, size=10.5, space_after=4)

# ── Signature block ──
add_divider(doc)
add_heading(doc, "SIGNATURES", 13, space_before=16, space_after=10)

sig_table = doc.add_table(rows=4, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT

sig_data = [
    ("Developer: Ian Kinoti",             "Client: [CLIENT_FULL_NAME]"),
    ("Signature: ____________________",   "Signature: ____________________"),
    ("Date: __________________________",  "Date: __________________________"),
    ("Email: kicoinkinoti@gmail.com",     "Email: [CLIENT_EMAIL]"),
]

for row_data in sig_data:
    row = sig_table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, row_data)):
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(10)
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = (row_data.index(text) == 0 and i < 1) or text.startswith("Client:")

# ── Footer note ──
add_paragraph(doc, "", space_after=10)
add_paragraph(doc,
    "By signing above, both parties confirm they have read, understood, and agreed to all terms in this document.",
    size=9, colour=MID_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=10)

footer_line = doc.add_paragraph()
footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(footer_line, "1A1A2E")
run = footer_line.add_run("  Ian Kinoti  ·  kicoinkinoti@gmail.com  ·  https://my-portfolio-ten-jet-99.vercel.app  ·  Kenya  ")
set_font(run, size=8, colour=WHITE)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
output_path = "/home/khelan/microsoft_word/Ian_Kinoti_Client_Document.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
